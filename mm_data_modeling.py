# -*- coding: utf-8 -*-
import os
from docx import Document
from datetime import datetime
import pandas as pd
import json
import googleAddressLocator as goog

punctuation = "!\"#$%&'()*+,-./:;<=>?@[\\]^_`{|}~"
def remove_punctuation(s):
    s_sans_punct = ""
    for letter in s:
        if letter not in punctuation:
            s_sans_punct += letter
        elif letter in punctuation:
            s_sans_punct += ' '
    return s_sans_punct

centroids = pd.read_csv('countrycentroidsmm.csv', usecols=['Country_name', 
    'UNc_latitude', 'UNc_longitude'])

# with open('countryTranslateDict.json') as json_file:
#     countryCrosswalk = json.load(json_file)

#stripping the country names in my centroid file of punctuation, 
#leading/trailing spaces, and converting to all lower-case.
centroids['Country_name'] = centroids['Country_name'].apply(lambda x: remove_punctuation(x))
centroids['Country_name'] = centroids['Country_name'].apply(lambda x: x.strip())
centroids['Country_name'] = centroids['Country_name'].apply(lambda x: x.lower())


dates = []
file_names = []
countries = []
regions = []
stories = []
titles = []
links = []
lat = []
lng = []
countries_simple = []

regions_codes = {'GENERAL':'General',\
                 'CEE/CIS':'Central and Eastern Europe and the Commonwealth of Independent States',\
                 'EAP':'East Asia and Pacific',\
                 'EAPR':'East Asia and Pacific',\
                 'ESA':'Eastern and Southern Africa',\
                 'LAC':'Latin America and the Caribbean',\
                 'MENA':'Middle East and North Africa',\
                 'ROSA':'South Asia',\
                 'WCA':'West and Central Africa',\
                 'SA':'South Asia'}

folder = 'OPSCEN Brief 2014/'
# folder = 'test/'

special_files = ['UNICEF OPSCEN Brief – 29 December 2014.docx','UNICEF OPSCEN Brief – 30 December 2014.docx','UNICEF OPSCEN Brief – 31 December 2014.docx']

flag_country = False
flag_title = False
flag_story = False
curr_story = ''
            
for file_name in os.listdir(folder):
    try:
        f = open(folder+file_name)
        document = Document(f)
        # for p in document.paragraphs:
            
        #     if p.text == '' or 'Disclaimer:' in p.text or 'The Brief is produced' in p.text:
        #             continue

        #     if p.text in regions_codes:
        #         curr_region = regions_codes[p.text]
        #         flag_country = True
        #         continue
                
        #     if 'http' in p.text:
        #         regions.append(curr_region)
        #         countries.append(curr_country)
        #         titles.append(curr_title)
        #         stories.append(curr_story)
        #         file_names.append(file_name)
        #         if file_name == 'UNICEF OPSCEN Brief – 29 December 2014.docx':
        #             date = datetime(year=2014, month=12, day=29)
        #         elif file_name == 'UNICEF OPSCEN Brief – 30 December 2014.docx': 
        #             date = datetime(year=2014, month=12, day=30)
        #         elif file_name == 'UNICEF OPSCEN Brief – 31 December 2014.docx':
        #             date = datetime(year=2014, month=12, day=31)
        #         else:
        #             date = document.core_properties.modified

        #         dates.append(date)
        #         links.append(p.text)

        #         formattedcountry = remove_punctuation(curr_country).strip().lower()
        #         # googleGeolocation = goog.address_locator(formattedcountry)
        #         # lat.append(googleGeolocation['lat'])
        #         # lng.append(googleGeolocation['lng'])

        #         # if 'across' in formattedcountry:
        #         #     countries_simple.append('regional or global')
        #         # else:
        #         #     countries_simple.append(formattedcountry)
        #         countries_simple.append(formattedcountry)


        #         flag_country = True
        #         flag_title = False
        #         curr_story = ''
        #         continue
                
        #     if flag_country:
        #         if len(p.text.split())< 4:
        #             curr_country = p.text
        #             flag_country = False
        #             flag_title = True
        #             continue
        #         else:
        #             flag_country = False
        #             flag_title = True

        #     if flag_title:
        #         curr_title = p.text
        #         flag_title = False
        #         flag_story = True
        #         continue
                
        #     if flag_story:
        #         curr_story = curr_story + p.text
        #         continue


    except:
        print "Could not process file: " + file_name

    else:
        for p in document.paragraphs:
            
            if p.text == '' or 'Disclaimer:' in p.text or 'The Brief is produced' in p.text:
                    continue

            if p.text in regions_codes:
                curr_region = regions_codes[p.text]
                flag_country = True
                continue
                
            if 'http' in p.text:
                regions.append(curr_region)
                countries.append(curr_country)
                titles.append(curr_title)
                stories.append(curr_story)
                file_names.append(file_name)
                if file_name == 'UNICEF OPSCEN Brief – 29 December 2014.docx':
                    date = datetime(year=2014, month=12, day=29)
                elif file_name == 'UNICEF OPSCEN Brief – 30 December 2014.docx': 
                    date = datetime(year=2014, month=12, day=30)
                elif file_name == 'UNICEF OPSCEN Brief – 31 December 2014.docx':
                    date = datetime(year=2014, month=12, day=31)
                else:
                    date = document.core_properties.modified

                dates.append(date)
                links.append(p.text)

                formattedcountry = remove_punctuation(curr_country).strip().lower()
                googleGeolocation = goog.address_locator(formattedcountry)
                lat.append(googleGeolocation['lat'])
                lng.append(googleGeolocation['lng'])

                # if 'across' in formattedcountry:
                #     countries_simple.append('regional or global')
                # else:
                #     countries_simple.append(formattedcountry)
                countries_simple.append(formattedcountry)


                flag_country = True
                flag_title = False
                curr_story = ''
                continue
                
            if flag_country:
                if len(p.text.split())< 4:
                    curr_country = p.text
                    flag_country = False
                    flag_title = True
                    continue
                else:
                    flag_country = False
                    flag_title = True

            if flag_title:
                curr_title = p.text
                flag_title = False
                flag_story = True
                continue
                
            if flag_story:
                curr_story = curr_story + p.text
                continue



    f.close()

print len(regions)
print len(stories)
print len(countries)
print len(titles)
print len(links)
print len(dates)
print len(countries_simple)

"""
Add in matching on formatted countries == centroid country names, inserting 
the appropriate lat and long into the lat and lng lists.
"""

df = pd.DataFrame({'region':regions, 'country':countries, 'title': titles , 'story':stories,\
                   'link': links,\
                   'file_name':file_names,\
                   'date':dates,
                   'lat': lat,
                   'lng': lng,
                   'Country_name': countries_simple})

# df['country'] = df['country'].apply(lambda x: remove_punctuation(x))
df['country'] = df['country'].apply(lambda x: x.strip())
# df['country'].apply(lambda x: x.lower())

# dfmerged = pd.merge(df, centroids, how='left', on='Country_name', left_index=False)

df = pd.merge(df, centroids, how='left', on='Country_name', left_index=False)

# df.link.unique()

# print df.country.unique()

# df.region.unique()

df.to_csv('OPSCENoutput.csv', encoding = 'utf-8')

# dfmerged.to_csv('OPSCENoutputMERGE.csv', encoding='utf-8')

# print df.head()

